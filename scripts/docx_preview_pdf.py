#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document


FONT_FILE = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
FONT_DIR = "/usr/share/fonts/opentype/noto"


def _build_html(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts = [
        "<style>",
        "@font-face { font-family: cjk; src: url('NotoSansCJK-Regular.ttc'); }",
        "body{font-family:cjk; font-size:11pt; line-height:1.45; color:#111;}",
        "h1,h2,h3{margin:0 0 10px 0;}",
        "p{margin:0 0 8px 0; white-space:pre-wrap;}",
        "table{border-collapse:collapse; width:100%; margin:12px 0;}",
        "td,th{border:1px solid #888; padding:6px; vertical-align:top;}",
        "</style><body>",
    ]
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        style = (para.style.name or "").lower() if para.style else ""
        if not text:
            parts.append("<p>&nbsp;</p>")
        elif "heading 1" in style:
            parts.append(f"<h1>{text}</h1>")
        elif "heading 2" in style:
            parts.append(f"<h2>{text}</h2>")
        elif "heading" in style:
            parts.append(f"<h3>{text}</h3>")
        else:
            parts.append(f"<p>{text}</p>")
    for table in doc.tables:
        parts.append("<table>")
        for row in table.rows:
            parts.append("<tr>")
            for cell in row.cells:
                cell_text = (cell.text or "").strip().replace("\n", "<br>")
                parts.append(f"<td>{cell_text}</td>")
            parts.append("</tr>")
        parts.append("</table>")
    parts.append("</body>")
    return "".join(parts)


def docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    html = _build_html(docx_path)
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    archive = fitz.Archive(FONT_DIR)
    page.insert_htmlbox(fitz.Rect(36, 36, 559, 806), html, archive=archive)
    doc.save(str(pdf_path))
    doc.close()
