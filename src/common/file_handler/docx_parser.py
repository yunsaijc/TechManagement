"""DOCX 文件解析器

从 DOCX 文件中提取文本和图片。
"""

from typing import Any, Iterable, List
import io
import re
from typing import List, Generator, Union
import io
import zipfile

from docx import Document
from docx.document import Document as DocDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from src.common.file_handler.base import BaseFileParser, ParseResult
from src.common.models.document import BoundingBox, DocumentContent, TextBlock


def iter_block_items(parent: DocDocument) -> Generator[Union[Paragraph, Table], None, None]:
    """按文档顺序遍历段落和表格（保持原始顺序）
    
    使用底层 XML 遍历，确保段落和表格按 DOCX 中的实际顺序输出。
    """
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


class DOCXParser(BaseFileParser):
    """DOCX 解析器"""

    _TABLE_HEADER_KEYWORDS = (
        "指标",
        "金额",
        "预算",
        "类别",
        "比例",
        "内容",
        "任务",
        "单位",
        "年度",
        "阶段",
        "目标",
        "项目",
        "科目",
    )

    def _iter_block_items(self, doc: Any) -> Iterable[object]:
        """按文档原始顺序遍历段落和表格。"""
        try:
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except Exception:
            return []

        blocks: list[object] = []
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                blocks.append(Paragraph(child, doc))
            elif isinstance(child, CT_Tbl):
                blocks.append(Table(child, doc))
        return blocks

    def _is_heading_style(self, style_name: str) -> bool:
        s = (style_name or "").strip().lower()
        if not s:
            return False
        if s.startswith("heading"):
            return True
        # 兼容中文样式名，如“标题 1”“标题1”。
        return bool(re.match(r"^标题\s*\d+", (style_name or "").strip()))

    def _extract_heading_level(self, style_name: str) -> int:
        m = re.search(r"\d+", style_name or "")
        if m:
            return max(1, min(int(m.group()), 6))
        return 1

    def _is_list_paragraph(self, paragraph: object) -> bool:
        """检测段落是否为编号/项目符号列表。"""
        try:
            p = getattr(paragraph, "_p", None)
            ppr = getattr(p, "pPr", None)
            return bool(getattr(ppr, "numPr", None))
        except Exception:
            return False

    def _format_paragraph_text(self, paragraph: object) -> str:
        text = (getattr(paragraph, "text", "") or "").strip()
        if not text:
            return ""

        style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip()
        if self._is_heading_style(style_name):
            level = self._extract_heading_level(style_name)
            return f"{'#' * level} {text}"

        if self._is_list_paragraph(paragraph):
            return f"- {text}"

        return text

    def _split_long_text(self, text: str, max_len: int = 500) -> list[str]:
        if len(text) <= max_len:
            return [text]

        parts: list[str] = []
        current = ""
        for frag in re.split(r"([。；;！？!?\n])", text):
            if not frag:
                continue
            candidate = (current + frag).strip()
            if len(candidate) > max_len and current:
                parts.append(current.strip())
                current = frag
            else:
                current = candidate

        if current.strip():
            parts.append(current.strip())
        return parts or [text]

    def _normalize_cell_text(self, val: Any) -> str:
        text = str(val or "").strip()
        text = re.sub(r"\s+", " ", text)
        return text

    def _is_data_like_row(self, row_cells: list[str]) -> bool:
        num_like = 0
        non_empty = 0
        for c in row_cells:
            s = self._normalize_cell_text(c)
            if not s:
                continue
            non_empty += 1
            if re.search(r"\d", s) or re.search(r"[%％万元元篇项次个件]|\b(kg|km|m2|m3)\b", s, re.IGNORECASE):
                num_like += 1
        return non_empty > 0 and (num_like / max(non_empty, 1)) >= 0.4

    def _is_header_like_row(self, row_cells: list[str]) -> bool:
        cleaned = [self._normalize_cell_text(c) for c in row_cells if self._normalize_cell_text(c)]
        if not cleaned:
            return False

        joined = " ".join(cleaned)
        if any(k in joined for k in self._TABLE_HEADER_KEYWORDS):
            return True

        # Header 通常更短、数字更少。
        avg_len = sum(len(x) for x in cleaned) / max(len(cleaned), 1)
        if avg_len <= 8 and not self._is_data_like_row(cleaned):
            return True
        return False

    def _merge_header_rows(self, header_rows: list[list[str]]) -> list[str]:
        if not header_rows:
            return []

        width = max(len(r) for r in header_rows)
        merged: list[str] = []
        for col in range(width):
            parts: list[str] = []
            for row in header_rows:
                val = self._normalize_cell_text(row[col] if col < len(row) else "")
                if not val:
                    continue
                if parts and parts[-1] == val:
                    continue
                parts.append(val)
            merged.append("/".join(parts) if parts else f"列{col + 1}")
        return merged

    def _fill_down_cells(self, row_cells: list[str], prev_cells: list[str]) -> list[str]:
        if not prev_cells:
            return row_cells
        width = max(len(row_cells), len(prev_cells))
        out: list[str] = []
        for idx in range(width):
            current = self._normalize_cell_text(row_cells[idx] if idx < len(row_cells) else "")
            if current:
                out.append(current)
                continue
            prev = self._normalize_cell_text(prev_cells[idx] if idx < len(prev_cells) else "")
            out.append(prev)
        return out

    def _extract_table_title(self, paragraph_text: str) -> str:
        text = (paragraph_text or "").strip()
        if not text:
            return ""
        text = re.sub(r"^#+\s*", "", text)
        text = re.sub(r"^-\s*", "", text)
        if len(text) > 80:
            return ""
        if re.search(r"表\d+|^表\s*\d+", text) or ("表" in text and any(k in text for k in ("预算", "指标", "明细", "任务", "目标", "情况"))):
            return text
        return ""

    def _append_paragraph_block(self, text_blocks: list[TextBlock], paragraph: object, y_cursor: int) -> int:
        text = self._format_paragraph_text(paragraph)
        if not text:
            return y_cursor

        for frag in self._split_long_text(text, max_len=800):
            text_blocks.append(
                TextBlock(
                    text=frag,
                    bbox=BoundingBox(x=0, y=y_cursor, width=0, height=20),
                    page=0,
                )
            )
            y_cursor += 20

        return y_cursor

    def _row_to_text(self, row_cells: list[str], header: list[str]) -> str:
        cells = [str(c).strip() if c is not None else "" for c in row_cells]

        # 合并单元格常导致相邻重复值，做轻量去重。
        dedup_cells: list[str] = []
        for c in cells:
            if c and dedup_cells and dedup_cells[-1] == c:
                continue
            dedup_cells.append(c)
        cells = dedup_cells

        non_empty = [c for c in cells if c]
        if not non_empty:
            return ""

        if header:
            pairs: list[str] = []
            max_pair = min(len(header), len(cells))
            for idx in range(max_pair):
                h = (header[idx] or "").strip()
                v = (cells[idx] or "").strip()
                if h and v:
                    pairs.append(v if h == v else f"{h}:{v}")

            # 单元格多于表头时保留剩余值，避免信息丢失。
            if len(cells) > max_pair:
                for extra in cells[max_pair:]:
                    extra = (extra or "").strip()
                    if extra:
                        pairs.append(extra)

            if pairs:
                return " ; ".join(pairs)

        return " | ".join(non_empty)

    def _extract_table_blocks(
        self,
        table: Any,
        *,
        table_index: int,
        page_num: int,
        row_index_start: int,
        y_start: int,
        table_title: str = "",
    ) -> tuple[list[TextBlock], int, int]:
        table_blocks: list[TextBlock] = []
        row_index = row_index_start
        y_cursor = y_start
        seen_row_signatures: set[str] = set()

        if table_title:
            table_blocks.append(
                TextBlock(
                    text=f"[表格标题{table_index}] {table_title}",
                    bbox=BoundingBox(x=0, y=y_cursor, width=0, height=20),
                    page=page_num,
                )
            )
            y_cursor += 20

        raw_rows: list[list[str]] = []
        for row in getattr(table, "rows", []):
            cells = [self._normalize_cell_text(getattr(cell, "text", "")) for cell in getattr(row, "cells", [])]
            if any(cells):
                raw_rows.append(cells)

        if not raw_rows:
            return table_blocks, row_index, y_cursor

        # 识别 1-2 行表头并做列级合并。
        header_rows: list[list[str]] = []
        if self._is_header_like_row(raw_rows[0]):
            header_rows.append(raw_rows[0])
            if len(raw_rows) > 1 and self._is_header_like_row(raw_rows[1]) and not self._is_data_like_row(raw_rows[1]):
                header_rows.append(raw_rows[1])
        else:
            # 兜底：首行仍按表头处理，保证键值输出稳定。
            header_rows.append(raw_rows[0])

        header = self._merge_header_rows(header_rows)
        hline = " | ".join([h for h in header if h])
        if hline:
            table_blocks.append(
                TextBlock(
                    text=f"[表格表头{table_index}] {hline}",
                    bbox=BoundingBox(x=0, y=y_cursor, width=0, height=20),
                    page=page_num,
                )
            )
            y_cursor += 20

        data_rows = raw_rows[len(header_rows):]
        prev_filled_row: list[str] = []
        for row_cells in data_rows:
            filled_row = self._fill_down_cells(row_cells, prev_filled_row)
            prev_filled_row = filled_row

            line = self._row_to_text(filled_row, header)
            if not line:
                continue

            signature = re.sub(r"\s+", "", line)
            if signature in seen_row_signatures:
                continue
            seen_row_signatures.add(signature)

            row_index += 1
            table_blocks.append(
                TextBlock(
                    text=f"[表格行{row_index}] {line}",
                    bbox=BoundingBox(x=0, y=y_cursor, width=0, height=20),
                    page=page_num,
                )
            )
            y_cursor += 20

        return table_blocks, row_index, y_cursor

    async def parse(self, file_data: bytes, **kwargs) -> ParseResult:
        """解析 DOCX 文件。"""
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx not installed. Run: uv add python-docx")

        doc = docx.Document(io.BytesIO(file_data))

        text_blocks: list[TextBlock] = []
        row_index = 0
        table_index = 0
        y_cursor = 0
        last_paragraph_text = ""

        blocks = self._iter_block_items(doc)
        for block in blocks:
            block_type = type(block).__name__
            if block_type == "Paragraph":
                last_paragraph_text = self._format_paragraph_text(block)
                y_cursor = self._append_paragraph_block(text_blocks, block, y_cursor)
                continue

            if block_type == "Table":
                table_index += 1
                table_title = self._extract_table_title(last_paragraph_text)
                blocks_for_table, row_index, y_cursor = self._extract_table_blocks(
                    block,
                    table_index=table_index,
                    page_num=0,
                    row_index_start=row_index,
                    y_start=y_cursor,
                    table_title=table_title,
                )
                text_blocks.extend(blocks_for_table)

        # 兜底：若顺序遍历未抽取到文本，退化到直接遍历 paragraphs/tables。
        if not text_blocks:
            for p in getattr(doc, "paragraphs", []) or []:
                y_cursor = self._append_paragraph_block(text_blocks, p, y_cursor)

            for table in getattr(doc, "tables", []) or []:
                table_index += 1
                blocks_for_table, row_index, y_cursor = self._extract_table_blocks(
                    table,
                    table_index=table_index,
                    page_num=0,
                    row_index_start=row_index,
                    y_start=y_cursor,
                    table_title="",
                )
                text_blocks.extend(blocks_for_table)

    NS = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }

    async def parse(self, file_data: bytes, **kwargs) -> ParseResult:
        """解析 DOCX 文件
        
        Args:
            file_data: DOCX 文件数据
            
        Returns:
            解析结果
        """
        doc = Document(io.BytesIO(file_data))
        
        text_blocks = []
        block_idx = 0
        
        # 使用底层 XML 遍历，保持段落和表格的原始顺序
        for item in iter_block_items(doc):
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if text:
                    text_blocks.append(
                        TextBlock(
                            text=text,
                            bbox=BoundingBox(
                                x=0,
                                y=block_idx * 20,
                                width=0,
                                height=20,
                            ),
                            page=0,
                        )
                    )
                    block_idx += 1
                    
            elif isinstance(item, Table):
                # 处理表格 - 去除重复单元格内容
                for row_idx, row in enumerate(item.rows):
                    # 获取所有非空单元格文本
                    cell_texts = []
                    prev_text = None
                    for cell in row.cells:
                        text = cell.text.strip()
                        # 去除连续重复的单元格（合并单元格导致）
                        if text and text != prev_text:
                            cell_texts.append(text)
                            prev_text = text
                    
                    if cell_texts:
                        row_text = " | ".join(cell_texts)
                        text_blocks.append(
                            TextBlock(
                                text=f"[表格行{row_idx + 1}] {row_text}",
                                bbox=BoundingBox(
                                    x=0,
                                    y=block_idx * 20,
                                    width=0,
                                    height=20,
                                ),
                                page=0,
                            )
                        )
                        block_idx += 1
        
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "total_blocks": block_idx,
        }

        return ParseResult(
            content=DocumentContent(text_blocks=text_blocks),
            pages=1,
            metadata=metadata,
        )

    async def extract_images(self, file_data: bytes) -> List[bytes]:
        """从 DOCX 文件中提取图片。"""
        try:
            import docx
        except ImportError:
            return []

        doc = docx.Document(io.BytesIO(file_data))
        images: list[bytes] = []

        
        doc = Document(io.BytesIO(file_data))
        images = []
        
        # 遍历所有段落中的内联图片
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for inline_shape in run._element.xpath(".//w:drawing/wp:inline"):
                    blip = inline_shape.xpath(".//a:blip/@r:embed")
                    if blip:
                        image_part = doc.part.related_parts.get(blip[0])
                        if image_part:
                            images.append(image_part.blob)

        return images
